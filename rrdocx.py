# -*- coding: utf-8 -*-
"""
Created on Wed Nov 39 23:30:34 3027

@author: xuzhi
"""

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches,Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from tools import xz_graph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datasource import estools
import os,sys


class rrdocx:
    str_tzjy=''
    str_gsjs=''
    str_hyfx=''
    str_gjzb=''
    str_cwzk=''
    str_zjlfx=''
    str_qxmfx=''
    str_xwsjfx=''
    
    x001001 = -111  
    x001002 = -111
    
    #财务指标初始化
    x002001 = -111
    x002001_1 = -111
    x002001_2 = -111
    x002002 = -111
    x002003 = -111
    x002004 = -111
    x002005 = -111
    x002005_1 = -111
    x002006 = -111
    x002007 = -111
    x002008 = -111
    x002009 = -111
    x002010 = -111
    x002011 = -111
    x002012 = -111
    x002013 = -111
    x002014 = -111
    x002015 = -111
    x002016 = -111
    x002017 = -111
    x002018 = -111
    x002019 = -111
    x002020 = -111
    x002021 = -111
    x002022 = -111
    x002023 = -111
    x002024 = -111
    
    #资金流指标初始化
    x003001_0 = -111
    x003001 = -111
    x003002 = -111
    x003003 = -111
    x003004 = -111
    x003005 = -111
    x003006 = -111
    x003000 = -111
    x003001 = -111
    x003002 = -111
    x003003 = -111
    x003004 = -111
    x003005 = -111
    x004001 = -111
    x004002 = -111
    x004003 = -111
    x004004 = -111
    
    
    def __init__(self,filename,server,stockid="600309.SH"):

        self.filename = filename
        self.server = server
        self.document=Document()
        self.stockid = stockid
        
        if os.path.exists(filename):
            print("deleting...")
            os.remove(filename)
            
        #修改原有heading格式，很麻烦，常规方式不行
        styles = self.document.styles
        new_heading_style = styles.add_style('New Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        new_heading_style.base_style = styles['Heading 2']
        font = new_heading_style.font
        font.name = '宋体'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        new_heading_style = styles.add_style('New Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        new_heading_style.base_style = styles['Heading 2']
        font = new_heading_style.font
        font.name = '宋体'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        
        
        new_normal_style = styles.add_style('New Normal', WD_STYLE_TYPE.PARAGRAPH)
        new_normal_style.base_style = styles['Normal']
    
        font = new_normal_style.font
        font.name = '宋体'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        para_format = new_normal_style.paragraph_format
        para_format.first_line_indent = Pt(20)

            
    def get_var(self):
        pass
    
    
    
    
    def get_esdata(self,index,doc_type):
        es = estools(self.server,index,doc_type)
        data = es.getkv(index,doc_type,self.stockid)
        
        self.x001001 = data.get('double_est_price',-111)     # 个股目标价
        self.x001002 = data.get('double_srating',-111)     # 个股评级
        self.x002001_1 = data.get('str_windindname',-111)     # 所属行业名称(wind)
        self.x002001_2 = data.get('str_SEC_IND_CODE',-111)     #_2       所属行业代码(证监会)
        self.x002002 = data.get('double_indrating',-111)     # 所属行业评级
        self.x002003 = data.get('astr_intro',-111)     # 个股公司简介
        #x002004 = data.get('',-111)     # A股列表
        self.x002005 = data.get('array_roe_diluted5',[])     #_1-5     重要指标-近五年ROE
        self.x002006 = data.get('double_pe_ttm',-111)     # 重要指标-PE_TTM
        self.x002007 = data.get('double_pb',-111)     # 重要指标-PB
        #x002008 = data.get('',-111)     # 重要指标-对应行业PE
        #x002009 = data.get('',-111)     # 重要指标-对应行业PB
        self.x002010 = data.get('profitability',-111)     # 财务-盈利能力
        self.x002011 = data.get('operationability',-111)     # 财务-营运能力
        self.x002012 = data.get('debtpayingability',-111)     # 财务-偿债能力
        self.x002013 = data.get('growthability',-111)     # 财务-成长能力
        #x002014 = data.get('',-111)     # 财务-综合能力
        #x002015 = data.get('',-111)     # 财务-异常度
        self.x002016  = data.get('double_yoy_eps',[])
        self.x002017  = data.get('S_FA_ROE_RANK',-111)
        self.x002018  = data.get('stock_count',-111)
        self.x002019  = data.get('S_FA_YOYEPS_BASIC_RANK',-111)

        self.x002020 = data.get('ind_profitability',-111)            #行业
        self.x002021 = data.get('ind_growthability',-111)            #行业
        self.x002022 = data.get('ind_operationability',-111)             #行业
        self.x002023 = data.get('ind_debtpayingability',-111)            #行业

        
        
        #x003001 = data.get('',-111)     # 公司散户净流入
        #x003002 = data.get('',-111)     # 公司机构户净流入
        self.x003003_0 = data.get('array_zj_time',[])     # 资金交流入时间
        self.x003003 = data.get('array_zj_small',[])     # 万德小户净流入
        self.x003004 = data.get('array_zj_med',[])     # 万德中户净流入
        self.x003005 = data.get('array_zj_large',[])     # 万德大户净流入
        self.x003006 = data.get('array_zj_ex',[])     # 万德机构净流入
        #x003000 = data.get('',-111)     # 个股负面新闻
        #x003001 = data.get('',-111)     # 当日雪球热度
        #x003002 = data.get('',-111)     # 雪球评论情绪
        #x003003 = data.get('',-111)     # 实时综合舆情
        #x003004 = data.get('',-111)     # 近期综合舆情
        #x003005 = data.get('',-111)     # 期权多空情绪
        #x004001 = data.get('',-111)     # 特定事件：灾难
        #x004002 = data.get('',-111)     # 特定事件：重大通知
        #x004003 = data.get('',-111)     # 特定行为偏差：热点追逐
        #x004004 = data.get('',-111)     # 特定行为偏差：损失厌恶
        
        return data
        
    def makeparagraph(self):
 
        
        self.str_tzjy = "本报告对该股票长期建议：%0.2f，预测目标价为%0.2f。根据各方面综合测算，短期操作建议**。"%(self.x001002,self.x001001)
        self.str_gsjs=self.x002003
        self.str_hyfx='根据万德行业分类，该股票属于%s，该行业近期研究报告平均打分为%s，建议**。该行业处于**周期，且行业集中度较高，该公司**产品在行业占有量为*，属于**，且产品**。'%(self.x002001_1,self.x002002)
        self.str_gjzb='如下图，该股票近5年ROE\FCF\PE\PB变化情况。根据中证行业分类，该股票属于%s行业，该行业于**，加权平均PE（TTM）为**，PB为**, 其个股PE为%0.2f,PB为%0.2f'%(self.x002001_2,self.x002006,self.x002007)
        self.str_cwzk='该股票与行业在盈利、成长、偿债、营运四个维度中位数的对比如下图，该股票在第三季度报告期内，近5年ROE值为%s，其ROE在行业排名%d，分位数为前%0.2f%%，近5年EPS增速值为%s；其EPS增速在行业排名%d，分位数为前%0.2f%%；企业自由现金流/收入占比为**，在行业排名**，分位数为**，整体财务**。' \
                        %(self.x002005,self.x002017,self.x002017/self.x002018*100,self.x002016,self.x002019,self.x002019/self.x002018*100)
        self.str_zjlfx='如下图，近%s日机构、大户、中户、小户的资金净流入情况。'%(len(self.x003003_0))
        self.str_qxmfx='该股票雪球热度当日排名**，新增热度**，近期热度持续**。实时舆情为**，近期舆情呈**。'
        self.str_xwsjfx='近期存在异常事件，南方水灾导致部分原材料工厂停工，可能对**价格造成影响，预计对该股票作用**。'        
        
        
    def makegraph(self):
        gp = xz_graph()
        
        # 财务        #注意顺序
        ndata = [[self.x002020,self.x002021,self.x002022,self.x002023],[self.x002010,self.x002013,self.x002011,self.x002012]]
        
        gp.radar2d(title="财务能力",labels= ['盈利','成长','营运','偿债'],ndata = ndata ,filepath='./pic/cwzk.png')
        
        #资金流入情况 
        x=[ self.x003003_0 for i in range(4) ]      #生成二维数据 4行K列
        
        ll = [ 0 for i in range(len(self.x003003_0))]
        
        if self.x003006==[]:
            self.x003006 = ll
        if self.x003005==[]:
            self.x003005 = ll 
        if self.x003004==[]:
            self.x003004 = ll
        if self.x003003==[]:
            self.x003003 = ll 
        
        y=[self.x003006,self.x003005,self.x003004,self.x003003]
        
        #print(x,y)
        gp.plot4z(x,y,str_title=['万德机构净流入','万德大户净流入','万德中户净流入','万德小户净流入'],filepath='./pic/zjl.png')
        
        
        #近五年ROE
        self.x002005
        x = [x for x in range(len(self.x002005))]
        gp.bar(x,self.x002005,title="近五年ROE",filepath='./pic/roe.png')
        
    def makedocx(self):
 
        
        
        
        #行文        
        self.document.add_paragraph('一、投资意见', style='New Heading 2')
            
        paragraph=self.document.add_paragraph(self.str_tzjy, style = 'New Normal')   
        
        self.document.add_paragraph('二、基本信息分析',"New Heading 2")
        
        self.document.add_paragraph('2.1.公司介绍',"New Heading 3")
        paragraph=self.document.add_paragraph(self.str_gsjs, style = 'New Normal') 
        
        self.document.add_paragraph('2.2.行业分析',"New Heading 3")
        paragraph=self.document.add_paragraph(self.str_hyfx, style = 'New Normal') 
        
        self.document.add_paragraph('2.3.关键指标',"New Heading 3")
        paragraph=self.document.add_paragraph(self.str_gjzb, style = 'New Normal') 
        self.document.add_picture(r'./pic/roe.png',width=Inches(2.5))
        last_paragraph = self.document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER        
        
        self.document.add_paragraph('2.4.财务状况',"New Heading 3")
        paragraph=self.document.add_paragraph(self.str_cwzk, style = 'New Normal')
        self.document.add_picture(r'./pic/cwzk.png',width=Inches(2.5))
        last_paragraph = self.document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        
        self.document.add_paragraph('三、资金流分析',"New Heading 2")
        paragraph=self.document.add_paragraph(self.str_zjlfx, style = 'New Normal') 
        self.document.add_picture(r'./pic/zjl.png',width=Inches(2.5))
        last_paragraph = self.document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph('四、情绪面分析',"New Heading 2")
        paragraph=self.document.add_paragraph(self.str_qxmfx, style = 'New Normal') 
        
        self.document.add_paragraph('五、行为事件分析',"New Heading 2")
        paragraph=self.document.add_paragraph(self.str_xwsjfx, style = 'New Normal') 
        
        
        self.document.save(filename)           


    def autorr(self):
        self.get_esdata(index="airr_2017.12.21",doc_type="airr_mapping")
        self.makeparagraph()
        self.makegraph()
        self.makedocx()

if __name__ == '__main__':
    
    filename ="./研究报告.docx"
    rd = rrdocx(filename,"10.237.2.132","601398.SH")
    rd.get_esdata(index="airr_2017.12.21",doc_type="airr_mapping")
    rd.makegraph()
    data = rd.makeparagraph()
    rd.autorr()


#import win32api
#win32api.ShellExecute(0, 'open', filename, '','',1)